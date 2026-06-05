import io
import re
from datetime import datetime
from typing import Any

from fastapi import HTTPException
from fastapi.responses import StreamingResponse


def process_collab_export(body: dict[str, Any]) -> StreamingResponse:
    fmt = str(body.get("format") or "word").lower().strip()
    title = str(body.get("title") or "Documento").strip()
    html_content = str(body.get("html") or "").strip()
    text_content = str(body.get("text") or "").strip()

    if not html_content and not text_content:
        raise HTTPException(status_code=400, detail="Debes enviar el contenido del documento")

    if fmt == "excel":
        return _export_excel(title, text_content or _html_to_plain(html_content))
    return _export_word(title, html_content or text_content)


def _export_word(title: str, html_content: str) -> StreamingResponse:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph(f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # Parse HTML and write paragraphs with basic formatting
    for block in _parse_html_blocks(html_content):
        _write_block(doc, block)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{_safe(title)}_{_ts()}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_excel(title: str, text_content: str) -> StreamingResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    # Title row
    ws.merge_cells("A1:B1")
    cell = ws["A1"]
    cell.value = title
    cell.font = Font(bold=True, size=14, color="1F4E79")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 22

    # Date row
    ws.merge_cells("A2:B2")
    date_cell = ws["A2"]
    date_cell.value = f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    date_cell.font = Font(italic=True, size=9, color="888888")
    date_cell.alignment = Alignment(horizontal="right")

    # Content
    lines = [line.strip() for line in text_content.splitlines() if line.strip()]
    alt_fill = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")
    for row_idx, line in enumerate(lines, 4):
        c = ws.cell(row=row_idx, column=1, value=line)
        c.alignment = Alignment(wrap_text=True)
        if row_idx % 2 == 0:
            c.fill = alt_fill

    ws.column_dimensions["A"].width = 100

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"{_safe(title)}_{_ts()}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── HTML parsing ──────────────────────────────────────────────────────────────

def _parse_html_blocks(html: str) -> list[dict]:
    """Very lightweight HTML → block list converter for TipTap output."""
    blocks: list[dict] = []
    # Headings
    for level in range(1, 7):
        html = re.sub(
            rf"<h{level}[^>]*>(.*?)</h{level}>",
            lambda m, lv=level: f"__H{lv}__" + m.group(1) + "__END__",
            html,
            flags=re.DOTALL | re.IGNORECASE,
        )
    # Bold / italic inside paragraphs — keep markers for later
    html = re.sub(r"<strong[^>]*>(.*?)</strong>", r"**\1**", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<b[^>]*>(.*?)</b>", r"**\1**", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<em[^>]*>(.*?)</em>", r"_\1_", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<i[^>]*>(.*?)</i>", r"_\1_", html, flags=re.DOTALL | re.IGNORECASE)
    # List items
    html = re.sub(r"<li[^>]*>(.*?)</li>", r"• \1\n", html, flags=re.DOTALL | re.IGNORECASE)
    # Line breaks
    html = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    # Paragraphs → newline delimited
    html = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n", html, flags=re.DOTALL | re.IGNORECASE)
    # Strip remaining tags
    html = re.sub(r"<[^>]+>", "", html)

    for raw_line in html.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("__H") and "__END__" in line:
            level_str = line[3]
            text = line.split("__END__")[0][5:]
            text = re.sub(r"<[^>]+>", "", text).strip()
            if text:
                blocks.append({"type": "heading", "level": int(level_str), "text": text})
        else:
            text = re.sub(r"<[^>]+>", "", line).strip()
            if text:
                blocks.append({"type": "paragraph", "text": text})
    return blocks


def _write_block(doc: Any, block: dict) -> None:
    from docx.shared import Pt

    if block["type"] == "heading":
        level = min(block["level"], 9)
        doc.add_heading(block["text"], level=level)
    else:
        para = doc.add_paragraph()
        _apply_inline(para, block["text"])


def _apply_inline(para: Any, text: str) -> None:
    """Parse **bold** and _italic_ markers and add runs accordingly."""
    pattern = re.compile(r"(\*\*.*?\*\*|_.*?_)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _html_to_plain(html: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def _safe(text: str) -> str:
    return ("".join(c for c in text if c.isalnum() or c in " _-").strip() or "documento")[:40]


def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")
